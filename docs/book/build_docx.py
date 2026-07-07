#!/usr/bin/env python3
"""Convert xauby_book.html into a native, editable .docx.

Walks the same HTML source used for the PDF and rebuilds it with python-docx:
real Word heading styles (so References > Update Table lets the reader
regenerate a native TOC), tables, lists, code blocks, and note/warning callouts.
Thai text is forced onto a font with real Thai glyph coverage on both Windows
and Mac Word (Leelawadee UI / fallback Tahoma) via explicit east-Asian/complex-
script font run properties, since python-docx's high-level API only sets the
Latin font name.
"""
import re

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm, RGBColor, Twips

HTML_PATH = "xauby_book.html"
DOCX_PATH = "xAuby_Book_TH.docx"

THAI_FONT = "Leelawadee UI"
MONO_FONT = "Consolas"
GOLD = RGBColor(0xB8, 0x86, 0x0B)
INK = RGBColor(0x12, 0x20, 0x2C)
NOTE_BG = "FFF7E6"
WARN_BG = "FDECEB"
CODE_BG = "1C2530"
CODE_FG = RGBColor(0xEA, 0xE6, 0xDA)


def set_thai_font(run, size=None, bold=None, color=None, mono=False):
    font_name = MONO_FONT if mono else THAI_FONT
    run.font.name = font_name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(attr), font_name)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def shade_paragraph(paragraph, hex_color):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def add_left_border(paragraph, hex_color):
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "24")
    left.set(qn("w:space"), "4")
    left.set(qn("w:color"), hex_color)
    pbdr.append(left)
    pPr.append(pbdr)


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_inline(paragraph, node, base_bold=False, base_color=None):
    """Recursively render inline children of an HTML node into docx runs."""
    for child in node.children:
        if isinstance(child, NavigableString):
            text = str(child)
            if text.strip() == "" and "\n" in text:
                continue
            text = re.sub(r"\s+", " ", text)
            if not text:
                continue
            r = paragraph.add_run(text)
            set_thai_font(r, bold=base_bold, color=base_color)
        elif isinstance(child, Tag):
            if child.name == "strong":
                add_inline(paragraph, child, base_bold=True, base_color=INK)
            elif child.name in ("span",) and "mono" in (child.get("class") or []):
                r = paragraph.add_run(child.get_text())
                set_thai_font(r, mono=True, size=9.5, color=RGBColor(0x7A, 0x3B, 0x12))
            elif child.name == "code":
                r = paragraph.add_run(child.get_text())
                set_thai_font(r, mono=True, size=9.5, color=RGBColor(0x7A, 0x3B, 0x12))
            elif child.name == "a":
                add_inline(paragraph, child, base_bold=base_bold, base_color=RGBColor(0x14, 0x50, 0x8C))
            elif child.name == "em":
                r = paragraph.add_run(child.get_text())
                set_thai_font(r, bold=base_bold, color=base_color)
                r.italic = True
            else:
                add_inline(paragraph, child, base_bold=base_bold, base_color=base_color)


def add_para(doc, tag, style=None, align=None):
    p = doc.add_paragraph(style=style)
    if align:
        p.alignment = align
    add_inline(p, tag)
    if not tag.get_text(strip=True):
        p.add_run("")
    return p


def add_heading(doc, tag, level):
    text = tag.get_text(" ", strip=True)
    h = doc.add_heading(level=level)
    r = h.add_run(text)
    set_thai_font(r, size={1: 20, 2: 15, 3: 13, 4: 11.5}.get(level, 12), bold=True, color=INK)
    if level == 2:
        pPr = h._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "16")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), "B8860B")
        pbdr.append(bottom)
        pPr.append(pbdr)
    return h


def add_list(doc, tag, ordered=False, level=0):
    style = "List Number" if ordered else "List Bullet"
    for li in tag.find_all("li", recursive=False):
        # nested list?
        nested = li.find(["ul", "ol"], recursive=False)
        # text of this li excluding nested list text
        li_clone_text_node = BeautifulSoup(
            "".join(str(c) for c in li.children if not (isinstance(c, Tag) and c.name in ("ul", "ol"))),
            "lxml",
        )
        p = doc.add_paragraph(style=style)
        p.paragraph_format.left_indent = Cm(0.5 + 0.5 * level)
        add_inline(p, li_clone_text_node)
        if nested:
            add_list(doc, nested, ordered=nested.name == "ol", level=level + 1)


def add_table(doc, tag):
    rows = tag.find_all("tr")
    if not rows:
        return
    ncols = max(len(r.find_all(["th", "td"])) for r in rows)
    table = doc.add_table(rows=0, cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        cells_src = row.find_all(["th", "td"])
        is_header = row.find("th") is not None
        docx_row = table.add_row()
        for c_idx in range(ncols):
            cell = docx_row.cells[c_idx]
            cell.paragraphs[0].text = ""
            if c_idx < len(cells_src):
                src = cells_src[c_idx]
                p = cell.paragraphs[0]
                add_inline(p, src, base_bold=is_header, base_color=(RGBColor(0xF4, 0xEC, 0xD8) if is_header else None))
                colspan = src.get("colspan")
                if colspan and int(colspan) > 1 and c_idx + int(colspan) <= ncols:
                    for extra in range(1, int(colspan)):
                        cell.merge(docx_row.cells[c_idx + extra])
            if is_header:
                shade_cell(cell, "12202C")
            elif r_idx % 2 == 0:
                shade_cell(cell, "FAF8F2")
    doc.add_paragraph()


def add_pre(doc, tag):
    text = tag.get_text()
    lines = text.split("\n")
    p = doc.add_paragraph()
    shade_paragraph(p, CODE_BG)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.3)
    for i, line in enumerate(lines):
        if i > 0:
            p.add_run().add_break()
        r = p.add_run(line if line else " ")
        set_thai_font(r, mono=True, size=9, color=CODE_FG)


def add_callout(doc, tag, bg):
    p = doc.add_paragraph()
    shade_paragraph(p, bg)
    add_left_border(p, "B8860B" if bg == NOTE_BG else "C0392B")
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.2)
    add_inline(p, tag)


def render_node(doc, tag):
    name = tag.name
    if name == "h1":
        add_heading(doc, tag, 1)
    elif name == "h2":
        add_heading(doc, tag, 2)
    elif name == "h3":
        add_heading(doc, tag, 3)
    elif name == "h4":
        add_heading(doc, tag, 4)
    elif name == "p":
        cls = tag.get("class") or []
        if "chapter-kicker" in cls:
            p = doc.add_paragraph()
            r = p.add_run(tag.get_text(strip=True))
            set_thai_font(r, size=10.5, bold=True, color=GOLD)
        elif "small" in cls or "figcap" in cls:
            p = doc.add_paragraph()
            add_inline(p, tag)
            for r in p.runs:
                r.font.size = Pt(8.5)
        else:
            add_para(doc, tag)
    elif name == "ul":
        add_list(doc, tag, ordered=False)
    elif name == "ol":
        add_list(doc, tag, ordered=True)
    elif name == "table":
        add_table(doc, tag)
    elif name == "pre":
        add_pre(doc, tag)
    elif name == "div":
        cls = tag.get("class") or []
        if "note" in cls:
            add_callout(doc, tag, NOTE_BG)
        elif "warn" in cls:
            add_callout(doc, tag, WARN_BG)
        elif "chapter-kicker" in cls:
            p = doc.add_paragraph()
            r = p.add_run(tag.get_text(strip=True))
            set_thai_font(r, size=10.5, bold=True, color=GOLD)
        elif "grid2" in cls or cls == []:
            for child in tag.find_all(recursive=False):
                render_node(doc, child)
    elif name == "hr":
        doc.add_paragraph()
    # ignore anything else at block level


def add_titlepage(doc, cover):
    doc.add_paragraph()
    doc.add_paragraph()
    kicker = cover.select_one(".kicker")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(kicker.get_text(strip=True))
    set_thai_font(r, size=12, bold=True, color=GOLD)

    h1 = cover.find("h1")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h1.get_text(strip=True))
    set_thai_font(r, size=48, bold=True, color=INK)

    sub = cover.select_one(".sub")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(sub.get_text(strip=True))
    set_thai_font(r, size=14, color=RGBColor(0x33, 0x33, 0x33))

    doc.add_paragraph()
    meta = cover.select_one(".meta")
    for line in meta.get_text("\n", strip=True).split("\n"):
        if not line.strip():
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line.strip())
        set_thai_font(r, size=9.5, color=RGBColor(0x66, 0x66, 0x66))
    doc.add_page_break()


def add_native_toc(doc):
    p = doc.add_paragraph()
    r = p.add_run("สารบัญ")
    set_thai_font(r, size=20, bold=True, color=INK)
    hint = doc.add_paragraph()
    hr = hint.add_run(
        "(ตารางนี้สร้างจากหัวข้อจริงในเอกสาร — คลิกขวาแล้วเลือก ‘Update Field’ "
        "หรือกด F9 เพื่อดึงเลขหน้า/หัวข้อล่าสุดหลังแก้ไข)"
    )
    set_thai_font(hr, size=9, color=RGBColor(0x66, 0x66, 0x66))
    hr.italic = True

    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    fldChar3 = OxmlElement("w:t")
    fldChar3.text = "สารบัญจะปรากฏที่นี่หลังเปิดไฟล์ใน Word และ Update Field (F9)"
    fldChar4 = OxmlElement("w:fldChar")
    fldChar4.set(qn("w:fldCharType"), "end")
    r_element = run._r
    r_element.append(fldChar1)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar3)
    r_element.append(fldChar4)
    doc.add_page_break()


def setup_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = THAI_FONT
    normal.font.size = Pt(11)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:eastAsia", "w:cs"):
        rfonts.set(qn(attr), THAI_FONT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)


def main():
    global HTML_PATH, DOCX_PATH, THAI_FONT

    import argparse

    ap = argparse.ArgumentParser()
    ap.add_argument("--html", default=HTML_PATH)
    ap.add_argument("--docx", default=DOCX_PATH)
    ap.add_argument("--font", default=THAI_FONT,
                     help="Body font (rFonts ascii/eastAsia/cs); use a Thai-capable font for Thai content")
    args = ap.parse_args()
    HTML_PATH, DOCX_PATH, THAI_FONT = args.html, args.docx, args.font

    with open(HTML_PATH, encoding="utf-8") as f:
        soup = BeautifulSoup(f.read(), "lxml")

    doc = Document()
    setup_styles(doc)

    body = soup.find("body")
    cover = body.select_one(".cover")
    toc_page = body.select_one(".toc-page")

    add_titlepage(doc, cover)
    add_native_toc(doc)

    top_sections = [c for c in body.find_all(["div"], recursive=False)]
    for sec in top_sections:
        cls = sec.get("class") or []
        if "cover" in cls or "toc-page" in cls:
            continue
        for child in sec.find_all(recursive=False):
            render_node(doc, child)
        doc.add_page_break()

    doc.save(DOCX_PATH)
    print(f"wrote {DOCX_PATH}")


if __name__ == "__main__":
    main()
